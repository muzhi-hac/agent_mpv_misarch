#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT.parent / "outputs" / "misarch-agent-interaction-plan-20260610"
ASSETS = OUT / "assets"
DOCX_PATH = OUT / "MiSArch_Agent_Interaction_Refactoring_Plan_zh.docx"

FONT_CJK = "STHeiti"
FONT_LATIN = "Calibri"
PIL_FONT = Path("/System/Library/Fonts/STHeiti Medium.ttc")

NAVY = "16324F"
BLUE = "2E74B5"
SKY = "EAF3FA"
PALE_BLUE = "F2F7FB"
INK = "1F2933"
MUTED = "66727D"
LIGHT = "F2F4F7"
LINE = "CBD5DF"
GREEN = "2F7D5B"
PALE_GREEN = "EAF6F0"
GOLD = "9A6B00"
PALE_GOLD = "FFF6DB"
RED = "A33A3A"
PALE_RED = "FBECEC"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=10.5, bold=False, color=INK, italic=False, latin=FONT_LATIN) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def set_paragraph(
    paragraph,
    before=0,
    after=6,
    line=1.10,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    keep_with_next=False,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.alignment = align
    fmt.keep_with_next = keep_with_next


def add_text(
    doc: Document,
    text: str,
    size=10.5,
    bold=False,
    color=INK,
    before=0,
    after=6,
    line=1.10,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    italic=False,
    keep_with_next=False,
):
    p = doc.add_paragraph()
    set_paragraph(p, before, after, line, align, keep_with_next)
    r = p.add_run(text)
    set_font(r, size, bold, color, italic)
    return p


def add_rich_line(doc: Document, parts: Iterable[tuple[str, bool, str]], after=5, size=10.5):
    p = doc.add_paragraph()
    set_paragraph(p, after=after)
    for text, bold, color in parts:
        set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_bullet(doc: Document, text: str, level=0, color=INK, after=4):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_paragraph(p, after=after, line=1.12)
    set_font(p.add_run(text), size=10.3, color=color)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    set_paragraph(p, after=5, line=1.12)
    set_font(p.add_run(text), size=10.3)
    return p


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph(
        p,
        before={1: 12, 2: 9, 3: 7}[level],
        after={1: 7, 2: 5, 3: 4}[level],
        line=1.0,
        keep_with_next=True,
    )
    r = p.add_run(text)
    set_font(
        r,
        size={1: 16, 2: 13, 3: 11.5}[level],
        bold=True,
        color={1: NAVY, 2: BLUE, 3: NAVY}[level],
    )
    return p


def add_callout(doc: Document, label: str, text: str, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph(p, before=1, after=1, line=1.12)
    set_font(p.add_run(label + "  "), size=10.5, bold=True, color=accent)
    set_font(p.add_run(text), size=10.5, color=INK)
    add_text(doc, "", after=2)
    return table


def add_caption(doc: Document, text: str):
    return add_text(
        doc,
        text,
        size=8.8,
        color=MUTED,
        before=2,
        after=7,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
    )


def add_picture(doc: Document, path: Path, width=6.28):
    p = doc.add_paragraph()
    set_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(path), width=Inches(width))
    return p


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, left, first, after in (
        ("List Bullet", 0.50, -0.25, 4),
        ("List Bullet 2", 0.75, -0.25, 3),
        ("List Number", 0.50, -0.25, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(10.3)
        style.paragraph_format.left_indent = Inches(left)
        style.paragraph_format.first_line_indent = Inches(first)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.12


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_font(run, size=8.5, color=MUTED)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.38)

    header = section.header
    p = header.paragraphs[0]
    set_paragraph(p, after=0)
    set_font(p.add_run("MiSArch Agentic Interoperability"), size=8.5, bold=True, color=MUTED)
    set_font(p.add_run("  |  Refactoring Brief"), size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    set_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_font(p.add_run("TU Berlin · 2026-06-10  |  "), size=8.5, color=MUTED)
    add_page_number(p)


def font(size: int, bold=False):
    return ImageFont.truetype(str(PIL_FONT), size=size)


def text_bbox(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont):
    return draw.textbbox((0, 0), text, font=fnt)


def wrap_text(draw, text: str, fnt, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        if text_bbox(draw, candidate, fnt)[2] <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def centered_multiline(draw, box, text, fnt, fill, spacing=8):
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, fnt, x2 - x1 - 34)
    heights = [text_bbox(draw, line, fnt)[3] - text_bbox(draw, line, fnt)[1] for line in lines]
    total = sum(heights) + spacing * (len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line, height in zip(lines, heights):
        width = text_bbox(draw, line, fnt)[2]
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=fnt, fill=fill)
        y += height + spacing


def box(draw, xy, title, subtitle="", fill=WHITE, outline=LINE, title_color=NAVY):
    draw.rounded_rectangle(xy, radius=22, fill=f"#{fill}", outline=f"#{outline}", width=3)
    x1, y1, x2, y2 = xy
    if subtitle:
        centered_multiline(draw, (x1 + 10, y1 + 15, x2 - 10, y1 + 78), title, font(28, True), f"#{title_color}", 4)
        centered_multiline(draw, (x1 + 16, y1 + 76, x2 - 16, y2 - 10), subtitle, font(21), f"#{INK}", 5)
    else:
        centered_multiline(draw, xy, title, font(25, True), f"#{title_color}", 5)


def arrow(draw, start, end, color=BLUE, width=6, label=""):
    draw.line([start, end], fill=f"#{color}", width=width)
    x2, y2 = end
    x1, y1 = start
    import math

    angle = math.atan2(y2 - y1, x2 - x1)
    length = 20
    for delta in (2.55, -2.55):
        point = (
            x2 + length * math.cos(angle + delta),
            y2 + length * math.sin(angle + delta),
        )
        draw.line([end, point], fill=f"#{color}", width=width)
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        bbox = text_bbox(draw, label, font(18, True))
        pad = 8
        draw.rounded_rectangle(
            (mx - bbox[2] / 2 - pad, my - 18, mx + bbox[2] / 2 + pad, my + 18),
            radius=8,
            fill="#FFFFFF",
        )
        draw.text((mx - bbox[2] / 2, my - 14), label, font=font(18, True), fill=f"#{color}")


def diagram_canvas(title: str, subtitle: str = "", size=(1800, 1050)):
    image = Image.new("RGB", size, "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.text((60, 42), title, font=font(40, True), fill=f"#{NAVY}")
    if subtitle:
        draw.text((60, 100), subtitle, font=font(23), fill=f"#{MUTED}")
    draw.line((60, 145, size[0] - 60, 145), fill=f"#{LINE}", width=3)
    return image, draw


def build_current_target_diagram(path: Path) -> None:
    image, draw = diagram_canvas("当前实现与目标实现", "关键差异不是有没有 LLM，而是有没有观察—决策—反馈循环")
    draw.rounded_rectangle((60, 190, 845, 965), radius=28, fill=f"#{PALE_RED}", outline=f"#{RED}", width=3)
    draw.rounded_rectangle((955, 190, 1740, 965), radius=28, fill=f"#{PALE_GREEN}", outline=f"#{GREEN}", width=3)
    draw.text((95, 220), "当前：固定工作流", font=font(32, True), fill=f"#{RED}")
    draw.text((990, 220), "目标：交互式 Agent", font=font(32, True), fill=f"#{GREEN}")

    left_boxes = [
        ((160, 320, 745, 425), "Prompt 已指定", "list_products → get_product"),
        ((160, 505, 745, 610), "程序替 Agent 决策", "自动选择 products[0]"),
        ((160, 690, 745, 795), "写操作直接执行", "没有等待用户确认"),
    ]
    for i, (xy, title, subtitle) in enumerate(left_boxes):
        box(draw, xy, title, subtitle, fill=WHITE, outline="E1A6A6", title_color=RED)
        if i < len(left_boxes) - 1:
            arrow(draw, ((xy[0] + xy[2]) // 2, xy[3] + 8), ((xy[0] + xy[2]) // 2, left_boxes[i + 1][0][1] - 10), color=RED)

    right_boxes = [
        ((1055, 300, 1640, 405), "Agent 读取目标", "必要时先向用户澄清"),
        ((1055, 455, 1640, 560), "Agent 动态调用工具", "根据 observation 选择下一步"),
        ((1055, 610, 1640, 715), "Agent 请求确认", "保存精确 pending action"),
        ((1055, 765, 1640, 870), "确认后执行", "返回订单草稿与完整 trace"),
    ]
    for i, (xy, title, subtitle) in enumerate(right_boxes):
        box(draw, xy, title, subtitle, fill=WHITE, outline="A9D7C1", title_color=GREEN)
        if i < len(right_boxes) - 1:
            arrow(draw, ((xy[0] + xy[2]) // 2, xy[3] + 6), ((xy[0] + xy[2]) // 2, right_boxes[i + 1][0][1] - 8), color=GREEN)
    image.save(path, quality=95)


def build_architecture_diagram(path: Path) -> None:
    image, draw = diagram_canvas("改造后的目标架构", "MCP Gateway 保持工具边界；Agent Orchestrator 新增目标理解、状态和决策责任")
    box(draw, (80, 430, 330, 650), "User / Chat UI", "自然语言需求\n确认 / 取消", fill=SKY, outline=BLUE, title_color=BLUE)
    draw.rounded_rectangle((430, 270, 920, 810), radius=22, fill=f"#{PALE_GREEN}", outline=f"#{GREEN}", width=3)
    centered_multiline(draw, (450, 282, 900, 330), "Agent Orchestrator", font(28, True), f"#{GREEN}")
    box(draw, (1030, 350, 1370, 730), "MCP Gateway", "tools/list\ntools/call\nschema + side effects", fill=PALE_BLUE, outline=BLUE, title_color=BLUE)
    box(draw, (1480, 430, 1730, 650), "MiSArch", "GraphQL Gateway\nCatalog / Order", fill=LIGHT, outline=NAVY, title_color=NAVY)
    arrow(draw, (330, 540), (430, 540), label="message")
    arrow(draw, (920, 480), (1030, 480), label="tool call")
    arrow(draw, (1030, 630), (920, 630), color=GREEN, label="observation")
    arrow(draw, (1370, 540), (1480, 540), color=NAVY, label="GraphQL")

    draw.rounded_rectangle((450, 330, 900, 425), radius=18, fill="#FFFFFF", outline=f"#{GREEN}", width=2)
    centered_multiline(draw, (450, 330, 900, 425), "1. Decide: ask / tool / final", font(23, True), f"#{GREEN}")
    draw.rounded_rectangle((450, 465, 900, 560), radius=18, fill="#FFFFFF", outline=f"#{GREEN}", width=2)
    centered_multiline(draw, (450, 465, 900, 560), "2. Observe tool result", font(23, True), f"#{GREEN}")
    draw.rounded_rectangle((450, 600, 900, 695), radius=18, fill="#FFFFFF", outline=f"#{GREEN}", width=2)
    centered_multiline(draw, (450, 600, 900, 695), "3. Update state and decide again", font(23, True), f"#{GREEN}")
    centered_multiline(draw, (465, 720, 885, 785), "session memory · policy · confirmation · trace", font(19, True), f"#{GREEN}")
    arrow(draw, (675, 425), (675, 465), color=GREEN)
    arrow(draw, (675, 560), (675, 600), color=GREEN)
    arrow(draw, (900, 648), (970, 648), color=GREEN)
    draw.arc((815, 300, 1010, 720), 275, 86, fill=f"#{GREEN}", width=5)
    image.save(path, quality=95)


def build_sequence_diagram(path: Path) -> None:
    image, draw = diagram_canvas("一次真正的多轮交互", "示例任务：找 30 欧以内的电子产品；合适的话购买 2 个")
    actors = [
        (155, "User"),
        (545, "Agent"),
        (960, "MCP Gateway"),
        (1435, "MiSArch"),
    ]
    for x, label in actors:
        draw.rounded_rectangle((x - 105, 175, x + 105, 245), radius=18, fill=f"#{SKY}", outline=f"#{BLUE}", width=3)
        centered_multiline(draw, (x - 105, 175, x + 105, 245), label, font(24, True), f"#{NAVY}")
        draw.line((x, 245, x, 990), fill=f"#{LINE}", width=3)

    events = [
        (290, 155, 545, "30€以内的电子产品，买2个"),
        (365, 545, 960, "tools/list"),
        (430, 960, 545, "available tools + schemas"),
        (500, 545, 960, "list_products(top_k=10)"),
        (565, 960, 1435, "GraphQL products"),
        (630, 1435, 960, "product observations"),
        (695, 960, 545, "Smart Plug · 27.99€"),
        (760, 545, 960, "get_product(selected_id)"),
        (825, 960, 545, "detail + read-only metadata"),
        (890, 545, 155, "是否创建 2 件 PENDING 订单？"),
        (950, 155, 545, "确认"),
    ]
    for y, x1, x2, label in events:
        color = GREEN if x2 < x1 else BLUE
        arrow(draw, (x1, y), (x2, y), color=color, width=4)
        bbox = text_bbox(draw, label, font(18, True))
        draw.rounded_rectangle(
            ((x1 + x2) / 2 - bbox[2] / 2 - 8, y - 31, (x1 + x2) / 2 + bbox[2] / 2 + 8, y - 5),
            radius=6,
            fill="#FFFFFF",
        )
        draw.text(((x1 + x2) / 2 - bbox[2] / 2, y - 28), label, font=font(18, True), fill=f"#{color}")
    image.save(path, quality=95)


def build_state_diagram(path: Path) -> None:
    image, draw = diagram_canvas("Agent 状态机与写操作确认门", "只读工具可自动执行；任何有副作用的工具必须经过用户确认")
    states = [
        ((80, 380, 330, 535), "IDLE", "等待用户目标", SKY, BLUE),
        ((420, 240, 735, 395), "PLANNING", "澄清或选择工具", PALE_BLUE, BLUE),
        ((825, 240, 1140, 395), "OBSERVING", "记录工具结果", PALE_GREEN, GREEN),
        ((420, 650, 735, 805), "AWAIT_CONFIRM", "冻结 pending action", PALE_GOLD, GOLD),
        ((825, 650, 1140, 805), "EXECUTING", "确认后调用写工具", PALE_RED, RED),
        ((1325, 430, 1640, 585), "COMPLETED", "最终答案 + trace", LIGHT, NAVY),
    ]
    for xy, title, subtitle, fill, color in states:
        box(draw, xy, title, subtitle, fill=fill, outline=color, title_color=color)

    arrow(draw, (330, 455), (420, 320), label="user message")
    arrow(draw, (735, 320), (825, 320), label="read tool")
    arrow(draw, (980, 395), (650, 650), color=GOLD, label="write proposed")
    arrow(draw, (735, 727), (825, 727), color=RED, label="confirmed")
    arrow(draw, (1140, 727), (1480, 585), color=NAVY, label="success")
    arrow(draw, (1140, 320), (1325, 485), color=NAVY, label="final")
    arrow(draw, (825, 365), (735, 365), color=GREEN, label="next decision")
    arrow(draw, (575, 805), (215, 535), color=MUTED, label="cancel / timeout")
    draw.text((80, 900), "安全不变量：未经确认，create_pending_order 的执行次数必须始终为 0。", font=font(25, True), fill=f"#{RED}")
    image.save(path, quality=95)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], font_size=9.1):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths, indent_dxa=120)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for index, text in enumerate(headers):
        cell = hdr.cells[index]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        set_paragraph(p, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_font(p.add_run(text), size=font_size, bold=True, color=NAVY)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, text in enumerate(values):
            if row_index % 2:
                set_cell_shading(cells[index], "FAFBFC")
            p = cells[index].paragraphs[0]
            set_paragraph(p, after=0, line=1.08)
            set_font(p.add_run(text), size=font_size, color=INK)
    return table


def add_code_block(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    p = cell.paragraphs[0]
    set_paragraph(p, after=0, line=1.0)
    for index, line in enumerate(text.splitlines()):
        r = p.add_run(line)
        set_font(r, size=8.6, color=INK, latin="Menlo")
        r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
        if index < len(text.splitlines()) - 1:
            r.add_break()
    return table


def build_docx() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "current_target": ASSETS / "01_current_vs_target.png",
        "architecture": ASSETS / "02_target_architecture.png",
        "sequence": ASSETS / "03_agent_sequence.png",
        "state": ASSETS / "04_confirmation_state_machine.png",
    }
    build_current_target_diagram(diagrams["current_target"])
    build_architecture_diagram(diagrams["architecture"])
    build_sequence_diagram(diagrams["sequence"])
    build_state_diagram(diagrams["state"])

    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    # Page 1
    add_text(doc, "TECHNICAL REFACTORING BRIEF", size=10, bold=True, color=BLUE, after=10)
    add_text(
        doc,
        "MiSArch Agent 交互改造方案",
        size=25,
        bold=True,
        color=NAVY,
        after=4,
        line=1.0,
    )
    add_text(
        doc,
        "从 MCP Tool Gateway 走向可观察、可确认、可恢复的交互式 Agent",
        size=13.5,
        color=MUTED,
        after=13,
    )
    add_rich_line(
        doc,
        [
            ("项目：", True, NAVY),
            ("misarch-agent-gateway-go", False, INK),
            ("    日期：", True, NAVY),
            ("2026-06-10", False, INK),
        ],
        after=12,
    )
    add_callout(
        doc,
        "核心判断",
        "当前代码已经是合格的 agent-facing MCP gateway，但仍缺少真正的 Agent Orchestrator。"
        "评测脚本中的工具顺序和商品选择基本预设，因此不能充分证明 Agent 会根据环境反馈继续决策。",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    add_picture(doc, diagrams["current_target"], width=6.25)
    add_caption(doc, "图 1  当前固定工作流与目标交互式 Agent 的差异")
    add_heading(doc, "建议结论", 2)
    add_bullet(doc, "保留现有 Go MCP Gateway；它继续负责工具 schema、输入验证和 MiSArch GraphQL 适配。")
    add_bullet(doc, "新增 Agent Orchestrator；它负责目标理解、动态工具选择、session memory、确认和最终回答。")
    add_bullet(doc, "答辩重点展示真实 trace：user message → decision → tool call → observation → next decision。")

    # Page 2
    add_page_break(doc)
    add_heading(doc, "1. 当前项目为什么会被认为“缺乏 Agent 交互”", 1)
    add_text(
        doc,
        "问题不在于项目没有 LLM 或 MCP，而在于 Agent 的关键决策没有真正发生在运行时。",
        size=11,
        bold=True,
        color=NAVY,
        after=8,
    )
    diagnosis_rows = [
        [
            "MCP server",
            "internal/mcpserver/server.go",
            "注册 list_products、get_product、create_pending_order；自身不理解用户目标。",
            "这是 Tool Server，不是 Agent。",
        ],
        [
            "LLM prompt",
            "scripts/agent_gcp_smoke_test.py",
            "Prompt 已给出 list_products → get_product 的计划模板。",
            "模型主要在复述预设计划。",
        ],
        [
            "结果选择",
            "run_mcp_agent",
            "程序将占位符替换为 products[0].product_id。",
            "没有比较商品后再选择。",
        ],
        [
            "写操作",
            "run_mcp_pending_order_agent",
            "直接调用 create_pending_order。",
            "没有 ask-user / confirm / cancel 状态。",
        ],
    ]
    add_table(
        doc,
        ["位置", "当前代码", "当前行为", "Tutor 看到的含义"],
        diagnosis_rows,
        [1400, 2050, 3160, 2750],
        font_size=8.8,
    )
    add_text(doc, "", after=2)
    add_heading(doc, "判定 Agent 交互的最低标准", 2)
    add_number(doc, "不同用户目标应产生不同的工具路径，而不是始终执行同一套调用。")
    add_number(doc, "下一步行为必须依赖上一步 observation，例如价格、类别、found=false 或工具错误。")
    add_number(doc, "信息不足时，Agent 能暂停并向用户提问。")
    add_number(doc, "有副作用的操作必须等待用户明确确认，并允许取消。")
    add_number(doc, "每一步 decision、action 和 observation 都能通过 trace 被展示和评估。")
    add_callout(
        doc,
        "一句话",
        "“LLM 调用了工具”不自动等于“存在 Agent 交互”；真正的证据是多轮反馈改变了后续行为。",
        fill=PALE_RED,
        accent=RED,
    )

    # Page 3
    add_page_break(doc)
    add_heading(doc, "2. 改造后的目标架构", 1)
    add_picture(doc, diagrams["architecture"], width=6.28)
    add_caption(doc, "图 2  新增 Agent Orchestrator，但保留现有 MCP Gateway 的边界")
    architecture_rows = [
        ["Chat/API", "接收 message、session_id、confirm/cancel；返回状态、消息和 trace_id。"],
        ["Agent Orchestrator", "循环执行 decide → act → observe；决定 ask_user、call_tool 或 final。"],
        ["Session Store", "保存对话、工具 observation、selected_product 和 pending_action。"],
        ["Policy / Confirmation", "只读工具自动允许；写工具生成待确认动作，确认后才执行。"],
        ["MCP Client", "完成 initialize、tools/list、tools/call，并将错误作为 observation 返回。"],
        ["MCP Gateway", "继续负责 schema、validation、side_effects、GraphQL adapter。"],
    ]
    add_table(doc, ["组件", "明确责任"], architecture_rows, [2200, 7160], font_size=9.2)
    add_text(doc, "", after=2)
    add_callout(
        doc,
        "架构边界",
        "不要把业务选择逻辑塞进 MCP tool handler。Tool 负责可靠执行；Agent 负责目标、选择和对话。",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    # Page 4
    add_page_break(doc)
    add_heading(doc, "3. 一次完整运行应该如何交互", 1)
    add_picture(doc, diagrams["sequence"], width=6.28)
    add_caption(doc, "图 3  Agent 根据商品 observation 继续决策，并在写操作前返回用户")
    trace_rows = [
        ["1", "user_message", "找 30 欧以内的电子产品，合适的话购买 2 个", "任务进入 session"],
        ["2", "decision", "call list_products / search_products", "选择只读工具"],
        ["3", "observation", "Smart Plug，27.99 EUR，Electronics", "获得选择依据"],
        ["4", "decision", "call get_product(selected_id)", "验证详情"],
        ["5", "ask_user", "是否创建 2 件 PENDING 订单？", "状态=awaiting_confirmation"],
        ["6", "user_message", "确认", "匹配 pending action"],
        ["7", "tool_call", "create_pending_order(exact stored args)", "执行受控写操作"],
        ["8", "final", "订单草稿已创建，尚未支付", "状态=completed"],
    ]
    add_table(doc, ["Step", "事件", "可展示内容", "状态变化"], trace_rows, [680, 1550, 4300, 2830], font_size=8.7)

    # Page 5
    add_page_break(doc)
    add_heading(doc, "4. 代码应该如何改", 1)
    add_text(
        doc,
        "推荐在同一 Go 项目中新增 Agent 层，使最终演示不再依赖一次性评测脚本。"
        "下面是最小可交付范围；文件名可以按现有项目风格调整。",
        after=8,
    )
    code_rows = [
        ["新增", "internal/agent/types.go", "定义 Decision、ToolCall、Observation、PendingAction、SessionState。"],
        ["新增", "internal/agent/orchestrator.go", "实现有最大步数的 decide → act → observe 循环。"],
        ["新增", "internal/agent/model_client.go", "向模型发送 history + discovered tools；解析结构化 decision。"],
        ["新增", "internal/agent/session_store.go", "保存 session history、selected product、pending action 和 trace。"],
        ["新增", "internal/agent/policy.go", "识别只读/写工具；写工具转为 awaiting_confirmation。"],
        ["新增", "internal/mcpclient/client.go", "封装 initialize、tools/list、tools/call 和 MCP session。"],
        ["新增", "internal/httpserver/agent_handler.go", "增加 POST /agent/chat 与 GET /agent/traces/:id。"],
        ["修改", "cmd/server/main.go", "注入 model、store、MCP client 和 orchestrator。"],
        ["修改", "internal/mcpserver/server.go", "为工具提供机器可判定的 readOnly/destructive metadata。"],
        ["扩展", "catalog/order tools", "增加 search_products 或筛选参数；增加 checkout_context 只读工具。"],
    ]
    add_table(doc, ["类型", "建议文件", "改动"], code_rows, [820, 2940, 5600], font_size=8.5)
    add_heading(doc, "核心循环伪代码", 2)
    add_code_block(
        doc,
        """for step := 0; step < maxSteps; step++ {
    decision := model.Decide(session.History, discoveredTools)

    switch decision.Type {
    case "ask_user":
        return AwaitingUser(decision.Question)
    case "tool_call":
        if policy.RequiresConfirmation(decision.Tool) {
            session.PendingAction = Freeze(decision.Tool, decision.Arguments)
            return AwaitingConfirmation(session.PendingAction)
        }
        observation := mcp.CallTool(decision.Tool, decision.Arguments)
        session.Append(decision, observation) // 再次交给模型决策
    case "final":
        return Completed(decision.Message)
    }
}""",
    )
    add_callout(
        doc,
        "必须删除的固定逻辑",
        "不要要求 plan 必须同时包含 list_products 和 get_product；不要默认使用 products[0]；"
        "不要在未确认时直接调用 create_pending_order。",
        fill=PALE_RED,
        accent=RED,
    )

    # Page 6
    add_page_break(doc)
    add_heading(doc, "5. 写操作确认与失败恢复", 1)
    add_picture(doc, diagrams["state"], width=6.28)
    add_caption(doc, "图 4  确认门是体现 User ↔ Agent 交互最有说服力的部分")
    add_heading(doc, "实现规则", 2)
    add_bullet(doc, "Agent 提议写操作时，保存工具名、完整参数、side_effects 和参数摘要 hash。")
    add_bullet(doc, "用户说“确认”后，只能执行已经冻结的 pending action，不能让模型临时改参数。")
    add_bullet(doc, "用户取消或 session 超时后，清空 pending action；写工具调用次数保持为 0。")
    add_bullet(doc, "工具失败时，将 error 记录为 observation；Agent 可换参数、换工具或向用户解释。")
    add_bullet(doc, "设置 max_steps，例如 8；超过限制时返回可解释错误，避免无限循环。")
    add_heading(doc, "建议的 API 状态", 2)
    add_code_block(
        doc,
        """POST /agent/chat
{ "session_id": "demo-1", "message": "确认" }

HTTP 200
{
  "status": "completed",
  "message": "PENDING 订单草稿已创建，尚未支付。",
  "trace_id": "tr_20260610_001",
  "order": { "order_status": "PENDING", "quantity": 2 }
}""",
    )

    # Page 7
    add_page_break(doc)
    add_heading(doc, "6. 改造完成后的展示结果", 1)
    add_heading(doc, "建议现场演示对话", 2)
    transcript = [
        ("User", "帮我找一个 30 欧以内的电子产品，合适的话买 2 个。", BLUE),
        ("Agent", "我会先查看可用商品，并在创建订单前向你确认。", GREEN),
        ("Trace", "tools/list → list_products(top_k=10)", MUTED),
        ("Observation", "HomeLink Smart Plug Twin Pack · 27.99 EUR · Electronics & Gadgets", NAVY),
        ("Trace", "get_product(0418eaa4-…)", MUTED),
        ("Agent", "推荐 HomeLink Smart Plug，27.99 欧。创建 2 件 PENDING 订单吗？不会支付。", GREEN),
        ("User", "确认。", BLUE),
        ("Trace", "create_pending_order(frozen arguments)", RED),
        ("Agent", "订单草稿已创建，状态 PENDING；支付和最终下单尚未发生。", GREEN),
    ]
    for speaker, message, color in transcript:
        add_rich_line(
            doc,
            [(f"{speaker}: ", True, color), (message, False, INK)],
            after=4,
            size=9.8,
        )
    add_heading(doc, "改造前后可观察结果", 2)
    result_rows = [
        ["工具路径", "固定 list → detail", "随用户目标和 observation 改变"],
        ["商品选择", "默认 products[0]", "Agent 按预算、类别和详情选择"],
        ["澄清问题", "无", "缺参数时返回 awaiting_user"],
        ["写操作", "可直接调用", "必须 awaiting_confirmation → confirmed"],
        ["错误处理", "脚本失败", "错误成为 observation，可恢复或解释"],
        ["展示证据", "最终 JSON / latency", "逐步 trace + session state + 最终结果"],
    ]
    add_table(doc, ["维度", "当前", "改造后"], result_rows, [1800, 3000, 4560], font_size=9.0)
    add_heading(doc, "验收指标（目标值，不是当前实测结果）", 2)
    metric_rows = [
        ["Task success", "≥ 90%", "10 个不同 catalog / order 场景"],
        ["Confirmation compliance", "100%", "未确认时写工具调用必须为 0"],
        ["Tool selection accuracy", "≥ 90%", "选择与用户约束一致的工具"],
        ["Recovery rate", "≥ 80%", "UUID 错误、found=false、上游错误后恢复"],
        ["Trace completeness", "100%", "每次 decision / action / observation 可追踪"],
    ]
    add_table(doc, ["指标", "建议目标", "评估方式"], metric_rows, [2600, 1700, 5060], font_size=8.8)

    # Page 8
    add_page_break(doc)
    add_heading(doc, "7. 最小实施顺序与答辩重点", 1)
    phases = [
        ("第一步：把固定脚本变成动态 loop", "模型每次只决定下一步；tool result 必须回填 history。"),
        ("第二步：增加 session 与 trace", "同一个 session 可以跨多次 HTTP 请求继续执行。"),
        ("第三步：增加 confirmation gate", "没有确认时绝不调用 create_pending_order。"),
        ("第四步：补齐 agent 所需只读工具", "search_products / checkout_context，避免 Agent 绕过 MCP 直接查 GraphQL。"),
        ("第五步：加入失败案例评测", "模糊目标、无匹配商品、错误 UUID、取消订单、上游故障。"),
    ]
    for title, detail in phases:
        add_rich_line(doc, [(title, True, NAVY), ("  " + detail, False, INK)], after=7, size=10.2)
    add_heading(doc, "Tutor 面前最值得展示的三个证据", 2)
    add_bullet(doc, "同一句“找商品并购买”的任务产生至少两次模型决策，第二次决策引用第一次工具 observation。")
    add_bullet(doc, "Agent 在 create_pending_order 前主动停止并等待用户；取消时系统中没有新订单。")
    add_bullet(doc, "打开 trace，能够看到每一步的 user message、decision、tool arguments、observation 和状态变化。")
    add_callout(
        doc,
        "最终表述",
        "原项目证明了 MiSArch 可以通过 MCP 暴露 agent-facing tools；改造后进一步证明 Agent "
        "能够围绕用户目标进行多轮规划、观察、确认和受控执行。",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_heading(doc, "不必为了“更像 Agent”而做的事情", 2)
    add_bullet(doc, "不必立即引入 multi-agent；单 Agent 的反馈循环已经足以回应 tutor。")
    add_bullet(doc, "不必让模型直接生成任意 GraphQL；MCP 的价值正是限制和结构化工具表面。")
    add_bullet(doc, "不必把所有业务能力都暴露出去；小而明确的工具集合更容易解释安全边界。")
    add_text(
        doc,
        "建议最终提交材料：交互式 demo、一次成功 trace、一次取消 trace、一次错误恢复 trace，以及本 PDF 中的架构与状态机图。",
        size=10.5,
        bold=True,
        color=NAVY,
        before=8,
        after=0,
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
